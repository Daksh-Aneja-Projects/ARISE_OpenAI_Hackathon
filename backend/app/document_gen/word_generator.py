"""
Word Document Generator — generates comprehensive bid response documents using python-docx.
Consolidates all agent outputs into a formal, professional document.
"""

import os
from typing import Any, Dict, Optional
from datetime import datetime


def _unwrap_agent_output(output) -> dict:
    """Unwrap agent output envelope if needed.

    Agent outputs are stored in DB as {status, agent, result: {...}, timestamp}.
    But some code paths pass the inner result dict directly.
    This normalises both shapes to just the inner result dict.
    """
    if not output or not isinstance(output, dict):
        return {}
    # If it has a 'result' key with a dict value, it's a wrapped envelope
    if "result" in output and isinstance(output.get("result"), dict):
        # But only unwrap if it also has 'status' or 'agent' keys (confirms envelope)
        if "status" in output or "agent" in output:
            return output["result"]
    return output


def _add_narrative_paragraphs(doc, text: str):
    """Add narrative text as formatted paragraphs, handling markdown-like content."""
    if not text:
        return
    for para in text.split("\n"):
        stripped = para.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("• "):
            clean = stripped.lstrip("-•* ").replace("**", "")
            doc.add_paragraph(clean, style="List Bullet")
        elif stripped[0:3] in [
            "1. ",
            "2. ",
            "3. ",
            "4. ",
            "5. ",
            "6. ",
            "7. ",
            "8. ",
            "9. ",
        ]:
            clean = stripped[3:].replace("**", "")
            doc.add_paragraph(clean, style="List Number")
        else:
            clean = stripped.replace("**", "")
            doc.add_paragraph(clean)


def _add_table(doc, headers: list, rows: list):
    """Add a formatted table to the document."""
    from docx.shared import Pt

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val) if val else ""


def generate_sow(
    bid_data: Dict[str, Any],
    output_dir: str = "../knowledge_base",
    filename: Optional[str] = None,
) -> Dict[str, Any]:
    """Generate a comprehensive bid response document from all agent outputs."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return {"status": "error", "message": "python-docx not installed"}

    doc = Document()

    # Style defaults
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    client_name = bid_data.get("client_name", "") or bid_data.get("manifest", {}).get(
        "client", {}
    ).get("name", "Client")
    bid_ref = bid_data.get("bid_reference", "")
    contract_type = (bid_data.get("contract_type", "") or "").upper() or "ENGAGEMENT"
    industry = bid_data.get("client_industry", "") or bid_data.get("manifest", {}).get(
        "client", {}
    ).get("industry", "")
    products = bid_data.get("products", []) or bid_data.get("manifest", {}).get(
        "rfp", {}
    ).get("products", [])

    # ══════════════════════════════════════════════════════
    # TITLE PAGE
    # ══════════════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph("")
    title = doc.add_heading("Bid Response Document", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(f"{contract_type} Services for {client_name}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style = doc.styles["Subtitle"]

    meta_lines = [
        f"Reference: {bid_ref}",
        f"Client: {client_name}",
        f"Industry: {industry}" if industry else None,
        f"Products: {', '.join(products)}" if products else None,
        f"Prepared by: {bid_data.get('manifest', {}).get('bidder_profile', {}).get('name', '[Bidder]')}",
        f"Date: {datetime.now().strftime('%B %d, %Y')}",
        "Classification: Confidential",
    ]
    for line in meta_lines:
        if line:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # TABLE OF CONTENTS placeholder
    # ══════════════════════════════════════════════════════
    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph(
        "(Update this field after opening in Microsoft Word: References → Update Table)"
    )
    doc.add_page_break()

    section_num = 0

    # ══════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY (from intake)
    # ══════════════════════════════════════════════════════
    intake = _unwrap_agent_output(bid_data.get("intake_output"))
    section_num += 1
    doc.add_heading(f"{section_num}. Executive Summary", level=1)

    narrative = intake.get("narrative", "")
    exec_summary = intake.get("executive_summary", "")
    if exec_summary:
        doc.add_paragraph(exec_summary)
    elif narrative:
        _add_narrative_paragraphs(doc, narrative)
    else:
        bidder_name = (
            bid_data.get("manifest", {})
            .get("bidder_profile", {})
            .get("name", "[Bidder]")
        )
        doc.add_paragraph(
            f"This document presents {bidder_name}'s comprehensive response to the {contract_type} "
            f"requirements issued by {client_name}. Our proposal outlines a structured approach "
            f"to delivering measurable business value through our proven delivery methodology."
        )

    # Extracted fields
    extracted = intake.get("extracted_fields", {})
    if extracted:
        doc.add_heading(f"{section_num}.1 Key Requirements Identified", level=2)
        for field_key, field_val in extracted.items():
            if isinstance(field_val, dict):
                label = field_key.replace("_", " ").title()
                value = field_val.get("value", "")
                field_val.get("confidence", "")
                if value:
                    display = (
                        ", ".join(value) if isinstance(value, list) else str(value)
                    )
                    doc.add_paragraph(f"{label}: {display}", style="List Bullet")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # 2. STRATEGIC ASSESSMENT (from bid_no_bid)
    # ══════════════════════════════════════════════════════
    manifest = bid_data.get("manifest", {})
    bid_score = manifest.get("bid_score", {})
    section_num += 1
    if bid_score or bid_data.get("win_probability"):
        doc.add_heading(f"{section_num}. Strategic Assessment", level=1)
        rec = bid_data.get("bid_recommendation", "")
        wp = bid_data.get("win_probability")
        if rec:
            doc.add_paragraph(f"Recommendation: {rec}")
        if wp:
            doc.add_paragraph(f"Assessed Win Probability: {wp * 100:.0f}%")

        # Score dimensions
        dimensions = bid_score.get("dimensions", {})
        if dimensions:
            doc.add_heading(f"{section_num}.1 Assessment Dimensions", level=2)
            rows = []
            for dim_name, dim_data in dimensions.items():
                if isinstance(dim_data, dict):
                    rows.append(
                        [
                            dim_name.replace("_", " ").title(),
                            str(dim_data.get("score", "")),
                            dim_data.get("rationale", ""),
                        ]
                    )
            if rows:
                _add_table(doc, ["Dimension", "Score", "Rationale"], rows)
                doc.add_paragraph("")

        # Narrative
        score_narrative = bid_score.get("narrative", "")
        if score_narrative:
            _add_narrative_paragraphs(doc, score_narrative)
        doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # 3. SCOPE OF WORK (from scope_builder)
    # ══════════════════════════════════════════════════════
    scope = _unwrap_agent_output(bid_data.get("scope_output"))
    scope_pkg = scope.get("scope_package", {}) if isinstance(scope, dict) else {}
    section_num += 1
    doc.add_heading(f"{section_num}. Scope of Work", level=1)

    scope_narrative = scope.get("narrative", "")
    if scope_narrative:
        _add_narrative_paragraphs(doc, scope_narrative)

    # New scope_by_platform structure
    scope_by_platform = scope_pkg.get("scope_by_platform", [])
    if scope_by_platform:
        doc.add_heading(f"{section_num}.1 Scope by Platform", level=2)
        for plat in scope_by_platform:
            if isinstance(plat, dict):
                plat_name = plat.get("platform", "Platform")
                effort = plat.get("platform_effort_days", 0)
                doc.add_heading(f"{plat_name} ({effort} effort-days)", level=3)
                summary = plat.get("scope_summary", "")
                if summary:
                    doc.add_paragraph(summary)
                wps = plat.get("work_packages", [])
                if wps:
                    rows = []
                    for wp in wps:
                        if isinstance(wp, dict):
                            rows.append(
                                [
                                    wp.get("id", ""),
                                    wp.get("name", ""),
                                    wp.get("rfp_ref", ""),
                                    str(wp.get("effort_days", "")),
                                ]
                            )
                    if rows:
                        _add_table(
                            doc,
                            ["ID", "Work Package", "RFP Ref", "Effort (days)"],
                            rows,
                        )
                        doc.add_paragraph("")

    # Legacy WBS format
    wbs = scope_pkg.get("work_breakdown_structure", [])
    if wbs and not scope_by_platform:
        doc.add_heading(f"{section_num}.1 Work Breakdown Structure", level=2)
        for phase in wbs:
            phase_name = phase.get("phase", "Phase")
            doc.add_heading(phase_name, level=3)
            for wp in phase.get("work_packages", []):
                wp_id = wp.get("id", "")
                wp_name = wp.get("name", "")
                wp_desc = wp.get("description", "")
                doc.add_paragraph(f"{wp_id} — {wp_name}", style="List Number")
                if wp_desc:
                    doc.add_paragraph(wp_desc, style="List Bullet")

    # Scope boundaries
    in_scope = scope_pkg.get(
        "in_scope", scope_pkg.get("scope_matrix", {}).get("in_scope", [])
    )
    out_of_scope = scope_pkg.get(
        "out_of_scope", scope_pkg.get("scope_matrix", {}).get("out_of_scope", [])
    )
    assumptions = scope_pkg.get(
        "assumptions", scope_pkg.get("scope_matrix", {}).get("assumptions", [])
    )

    sub = 2
    if in_scope:
        doc.add_heading(f"{section_num}.{sub} In-Scope Items", level=2)
        for item in in_scope:
            doc.add_paragraph(str(item), style="List Bullet")
        sub += 1
    if out_of_scope:
        doc.add_heading(f"{section_num}.{sub} Exclusions", level=2)
        for item in out_of_scope:
            doc.add_paragraph(str(item), style="List Bullet")
        sub += 1
    if assumptions:
        doc.add_heading(f"{section_num}.{sub} Assumptions", level=2)
        for item in assumptions:
            doc.add_paragraph(str(item), style="List Bullet")
        sub += 1

    # Effort summary
    total_effort = scope_pkg.get("total_effort_days", 0)
    if total_effort:
        doc.add_heading(f"{section_num}.{sub} Effort Summary", level=2)
        doc.add_paragraph(f"Total estimated effort: {total_effort} days")
        doc.add_paragraph(f"Confidence: {scope_pkg.get('effort_confidence', 'Medium')}")
        sub += 1

    # Team model table
    team = scope_pkg.get("team_model", [])
    if team and isinstance(team, list):
        doc.add_heading(f"{section_num}.{sub} Proposed Team", level=2)
        rows = []
        for r in team:
            if isinstance(r, dict):
                rows.append(
                    [
                        r.get("role", ""),
                        str(r.get("count", "")),
                        r.get("location", ""),
                        r.get("platform", ""),
                    ]
                )
        if rows:
            _add_table(doc, ["Role", "Count", "Location", "Platform"], rows)

    timeline = scope_pkg.get("timeline_months")
    if timeline:
        sub += 1
        doc.add_heading(f"{section_num}.{sub} Estimated Timeline", level=2)
        doc.add_paragraph(f"Estimated duration: {timeline} months")
        transition = scope_pkg.get("transition_weeks", "")
        if transition:
            doc.add_paragraph(f"Transition period: {transition} weeks")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # 4. SOLUTION ARCHITECTURE (from solution_architect)
    # ══════════════════════════════════════════════════════
    solution = _unwrap_agent_output(bid_data.get("solution_output"))
    section_num += 1
    doc.add_heading(f"{section_num}. Solution Architecture", level=1)

    sol_narrative = solution.get("narrative", "")
    if sol_narrative:
        _add_narrative_paragraphs(doc, sol_narrative)

    if isinstance(solution, dict):
        for key in [
            "architecture_overview",
            "operating_model",
            "integration_landscape",
            "technology_stack",
            "deployment_model",
        ]:
            val = solution.get(key)
            if val and isinstance(val, str):
                doc.add_heading(key.replace("_", " ").title(), level=2)
                _add_narrative_paragraphs(doc, val)
            elif val and isinstance(val, dict):
                doc.add_heading(key.replace("_", " ").title(), level=2)
                for sub_k, sub_v in val.items():
                    if isinstance(sub_v, str):
                        doc.add_paragraph(
                            f"{sub_k.replace('_', ' ').title()}: {sub_v}",
                            style="List Bullet",
                        )
                    elif isinstance(sub_v, list):
                        doc.add_paragraph(
                            f"{sub_k.replace('_', ' ').title()}:", style="List Bullet"
                        )
                        for item in sub_v:
                            doc.add_paragraph(str(item), style="List Bullet 2")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # 5. COMPETITIVE POSITIONING (from competitive_intel)
    # ══════════════════════════════════════════════════════
    competitive = _unwrap_agent_output(bid_data.get("competitive_output"))
    section_num += 1
    doc.add_heading(f"{section_num}. Competitive Positioning & Win Strategy", level=1)

    comp_narrative = competitive.get("narrative", "")
    if comp_narrative:
        _add_narrative_paragraphs(doc, comp_narrative)

    if isinstance(competitive, dict):
        win_themes = competitive.get("win_themes") or competitive.get(
            "key_differentiators"
        )
        if win_themes:
            doc.add_heading(f"{section_num}.1 Key Differentiators", level=2)
            if isinstance(win_themes, list):
                for theme in win_themes:
                    if isinstance(theme, dict):
                        doc.add_paragraph(
                            f"{theme.get('theme', theme.get('name', ''))}: {theme.get('description', '')}",
                            style="List Bullet",
                        )
                    else:
                        doc.add_paragraph(str(theme), style="List Bullet")
            elif isinstance(win_themes, str):
                _add_narrative_paragraphs(doc, win_themes)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # 6. COMMERCIAL MODEL (from commercial_model)
    # ══════════════════════════════════════════════════════
    commercial = _unwrap_agent_output(bid_data.get("commercial_output"))
    section_num += 1
    doc.add_heading(f"{section_num}. Commercial Model", level=1)

    comm_narrative = commercial.get("narrative", "")
    if comm_narrative:
        _add_narrative_paragraphs(doc, comm_narrative)

    if isinstance(commercial, dict):
        # P&L Model
        pl_model = commercial.get("pl_model", {})
        if pl_model:
            doc.add_heading(f"{section_num}.1 Financial Summary", level=2)

            revenue = pl_model.get("revenue", {})
            if revenue:
                rows = [
                    [
                        k.replace("_", " ").title(),
                        f"${v:,.0f}" if isinstance(v, (int, float)) else str(v),
                    ]
                    for k, v in revenue.items()
                ]
                _add_table(doc, ["Revenue Item", "Value"], rows)
                doc.add_paragraph("")

            costs = pl_model.get("costs", {})
            if costs:
                doc.add_heading(f"{section_num}.2 Cost Structure", level=2)
                rows = [
                    [
                        k.replace("_", " ").title(),
                        f"${v:,.0f}" if isinstance(v, (int, float)) else str(v),
                    ]
                    for k, v in costs.items()
                ]
                _add_table(doc, ["Cost Item", "Value"], rows)
                doc.add_paragraph("")

            prof = pl_model.get("profitability", {})
            if prof:
                doc.add_heading(f"{section_num}.3 Profitability", level=2)
                for k, v in prof.items():
                    label = k.replace("_", " ").title()
                    if isinstance(v, (int, float)):
                        display = (
                            f"{v:.1f}%"
                            if "percent" in k or "margin" in k
                            else f"${v:,.0f}"
                        )
                    else:
                        display = str(v)
                    doc.add_paragraph(f"{label}: {display}", style="List Bullet")

        # Resource plan
        resource_plan = commercial.get("resource_plan", [])
        if resource_plan:
            doc.add_heading(f"{section_num}.4 Resource Plan", level=2)
            rows = []
            for r in resource_plan:
                if isinstance(r, dict):
                    rows.append(
                        [
                            r.get("role", ""),
                            r.get("location", ""),
                            str(r.get("count", "")),
                            f"M{r.get('start_month', 1)}",
                        ]
                    )
            if rows:
                _add_table(doc, ["Role", "Location", "FTEs", "Start"], rows)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # 7. COMPLIANCE & RISK (from compliance_risk)
    # ══════════════════════════════════════════════════════
    compliance = _unwrap_agent_output(bid_data.get("compliance_output"))
    section_num += 1
    doc.add_heading(f"{section_num}. Compliance & Risk Assessment", level=1)

    comp_risk_narrative = compliance.get("narrative", "")
    if comp_risk_narrative:
        _add_narrative_paragraphs(doc, comp_risk_narrative)

    if isinstance(compliance, dict):
        risk_data = compliance.get(
            "contractual_risks", compliance.get("risk_register", [])
        )
        if risk_data and isinstance(risk_data, list):
            doc.add_heading(f"{section_num}.1 Risk Register", level=2)
            rows = []
            for risk in risk_data:
                if isinstance(risk, dict):
                    rows.append(
                        [
                            risk.get("id", risk.get("clause", "")),
                            risk.get(
                                "description",
                                risk.get("risk", risk.get("clause_summary", "")),
                            ),
                            risk.get(
                                "severity",
                                risk.get("impact", risk.get("risk_level", "")),
                            ),
                            risk.get(
                                "mitigation", risk.get("negotiation_position", "")
                            ),
                        ]
                    )
            if rows:
                _add_table(doc, ["ID/Clause", "Risk", "Severity", "Mitigation"], rows)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # 8. AUTOMATION & AI OPPORTUNITIES (from automation_ai)
    # ══════════════════════════════════════════════════════
    automation = _unwrap_agent_output(bid_data.get("automation_ai_output"))
    section_num += 1
    doc.add_heading(f"{section_num}. Automation & AI Opportunities", level=1)

    auto_narrative = automation.get("narrative", "")
    if auto_narrative:
        _add_narrative_paragraphs(doc, auto_narrative)

    if isinstance(automation, dict):
        opportunities = automation.get("opportunities", [])
        if opportunities:
            doc.add_heading(f"{section_num}.1 Opportunity Catalogue", level=2)
            rows = []
            for opp in opportunities:
                if isinstance(opp, dict):
                    rows.append(
                        [
                            opp.get("id", ""),
                            opp.get("title", opp.get("name", "")),
                            opp.get("platform", ""),
                            opp.get("priority", ""),
                            opp.get(
                                "estimated_savings", opp.get("effort_reduction", "")
                            ),
                        ]
                    )
            if rows:
                _add_table(
                    doc,
                    ["ID", "Opportunity", "Platform", "Priority", "Est. Savings"],
                    rows,
                )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # 9. TRANSITION & CHANGE MANAGEMENT (from transition_change)
    # ══════════════════════════════════════════════════════
    transition = _unwrap_agent_output(bid_data.get("transition_change_output"))
    section_num += 1
    doc.add_heading(f"{section_num}. Transition & Change Management", level=1)

    trans_narrative = transition.get("narrative", "")
    if trans_narrative:
        _add_narrative_paragraphs(doc, trans_narrative)

    if isinstance(transition, dict):
        # Transition plan
        trans_plan = transition.get("transition_plan", {})
        if isinstance(trans_plan, dict):
            phases = trans_plan.get("phases", [])
            if phases:
                doc.add_heading(f"{section_num}.1 Transition Phases", level=2)
                rows = []
                for phase in phases:
                    if isinstance(phase, dict):
                        rows.append(
                            [
                                phase.get("phase", phase.get("name", "")),
                                phase.get("duration", ""),
                                phase.get("description", phase.get("activities", "")),
                            ]
                        )
                if rows:
                    _add_table(doc, ["Phase", "Duration", "Activities"], rows)
                    doc.add_paragraph("")

        # Change management approach
        change_mgmt = transition.get("change_management", "")
        if change_mgmt:
            doc.add_heading(f"{section_num}.2 Change Management Approach", level=2)
            if isinstance(change_mgmt, str):
                _add_narrative_paragraphs(doc, change_mgmt)
            elif isinstance(change_mgmt, dict):
                for k, v in change_mgmt.items():
                    if isinstance(v, str):
                        doc.add_paragraph(
                            f"{k.replace('_', ' ').title()}: {v}", style="List Bullet"
                        )

        # Governance model
        governance = transition.get(
            "governance_model", transition.get("governance", "")
        )
        if governance:
            doc.add_heading(f"{section_num}.3 Governance Model", level=2)
            if isinstance(governance, str):
                _add_narrative_paragraphs(doc, governance)
            elif isinstance(governance, dict):
                for k, v in governance.items():
                    if isinstance(v, str):
                        doc.add_paragraph(
                            f"{k.replace('_', ' ').title()}: {v}", style="List Bullet"
                        )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # 10. CLIENT INTELLIGENCE (from client_intelligence)
    # ══════════════════════════════════════════════════════
    client_intel = _unwrap_agent_output(bid_data.get("client_intel_output"))
    section_num += 1
    if client_intel:
        doc.add_heading(f"{section_num}. Client Intelligence", level=1)

        ci_narrative = client_intel.get("narrative", "")
        if ci_narrative:
            _add_narrative_paragraphs(doc, ci_narrative)

        if isinstance(client_intel, dict):
            # Client profile
            profile = client_intel.get("client_profile", {})
            if profile and isinstance(profile, dict):
                doc.add_heading(f"{section_num}.1 Client Profile", level=2)
                for k, v in profile.items():
                    if v and isinstance(v, str):
                        doc.add_paragraph(
                            f"{k.replace('_', ' ').title()}: {v}", style="List Bullet"
                        )
                    elif v and isinstance(v, list):
                        doc.add_paragraph(
                            f"{k.replace('_', ' ').title()}: {', '.join(str(x) for x in v)}",
                            style="List Bullet",
                        )

            # Win strategy
            win_strat = client_intel.get(
                "win_strategy", client_intel.get("engagement_strategy", "")
            )
            if win_strat:
                doc.add_heading(f"{section_num}.2 Engagement Strategy", level=2)
                if isinstance(win_strat, str):
                    _add_narrative_paragraphs(doc, win_strat)
                elif isinstance(win_strat, dict):
                    for k, v in win_strat.items():
                        if isinstance(v, str):
                            doc.add_paragraph(
                                f"{k.replace('_', ' ').title()}: {v}",
                                style="List Bullet",
                            )

        doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # 11. PROPOSAL NARRATIVE (from proposal_writer)
    # ══════════════════════════════════════════════════════
    proposal = _unwrap_agent_output(bid_data.get("proposal_output"))
    section_num += 1
    if proposal:
        doc.add_heading(f"{section_num}. Proposal Narrative", level=1)
        proposal_narrative = proposal.get("narrative", "")
        if proposal_narrative:
            _add_narrative_paragraphs(doc, proposal_narrative)

        if isinstance(proposal, dict):
            proposal_sections = proposal.get(
                "sections", proposal.get("proposal_sections", [])
            )
            if proposal_sections and isinstance(proposal_sections, list):
                for ps in proposal_sections:
                    if isinstance(ps, dict):
                        ps_title = ps.get("title", ps.get("heading", ""))
                        ps_content = ps.get("content", ps.get("body", ""))
                        if ps_title and ps_content:
                            doc.add_heading(ps_title, level=2)
                            _add_narrative_paragraphs(doc, ps_content)

        doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # 12. DISCOVERY & CLARIFICATIONS (from discovery)
    # ══════════════════════════════════════════════════════
    discovery = _unwrap_agent_output(bid_data.get("discovery_output"))
    section_num += 1
    if discovery:
        doc.add_heading(f"{section_num}. Discovery & Clarifications", level=1)
        disc_narrative = discovery.get("narrative", "")
        if disc_narrative:
            _add_narrative_paragraphs(doc, disc_narrative)

        disc_analysis = discovery.get("discovery_analysis", {})
        if isinstance(disc_analysis, dict):
            categories = disc_analysis.get("discovery_categories", [])
            if categories:
                doc.add_heading(f"{section_num}.1 Discovery Questions", level=2)
                for cat in categories:
                    if isinstance(cat, dict):
                        cat_name = cat.get("category", "General")
                        doc.add_heading(cat_name, level=3)
                        for q in cat.get("questions", []):
                            if isinstance(q, dict):
                                priority = q.get("priority", "")
                                question = q.get("question", "")
                                doc.add_paragraph(
                                    f"[{priority.upper()}] {question}",
                                    style="List Bullet",
                                )

            pre_meeting = disc_analysis.get("pre_meeting_requests", [])
            if pre_meeting:
                doc.add_heading(
                    f"{section_num}.2 Pre-Meeting Document Requests", level=2
                )
                for req in pre_meeting:
                    if isinstance(req, dict):
                        doc.add_paragraph(
                            f"{req.get('document', req.get('request', ''))}: {req.get('purpose', '')}",
                            style="List Bullet",
                        )
                    else:
                        doc.add_paragraph(str(req), style="List Bullet")

        doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # 10. QUALITY ASSURANCE (from qa)
    # ══════════════════════════════════════════════════════
    qa = _unwrap_agent_output(bid_data.get("qa_output"))
    section_num += 1
    if qa:
        doc.add_heading(f"{section_num}. Quality Assurance Review", level=1)
        qa_narrative = qa.get("narrative", "")
        if qa_narrative:
            _add_narrative_paragraphs(doc, qa_narrative)

        # Quality scores table
        qa_scores = qa.get("quality_scores", {})
        if qa_scores:
            doc.add_heading(f"{section_num}.1 Quality Scores", level=2)
            rows = [
                [
                    k.replace("_", " ").title(),
                    f"{v}/100" if isinstance(v, (int, float)) else str(v),
                ]
                for k, v in qa_scores.items()
            ]
            if rows:
                _add_table(doc, ["Dimension", "Score"], rows)

        doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # 11. OUTPUT GENERATOR EXECUTIVE SUMMARY (if available)
    # ══════════════════════════════════════════════════════
    output_gen = _unwrap_agent_output(bid_data.get("output_generator_output"))
    if isinstance(output_gen, dict):
        exec_summ = output_gen.get("executive_summary", "")
        sow_outline = output_gen.get("sow_outline", "")
        if exec_summ:
            section_num += 1
            doc.add_heading(f"{section_num}. Consolidated Executive Summary", level=1)
            _add_narrative_paragraphs(doc, exec_summ)
            doc.add_page_break()
        if sow_outline:
            section_num += 1
            doc.add_heading(f"{section_num}. Statement of Work Outline", level=1)
            _add_narrative_paragraphs(doc, sow_outline)
            doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # DATA ANALYST OUTPUT (if available)
    # ══════════════════════════════════════════════════════
    data_analyst = _unwrap_agent_output(bid_data.get("data_analyst_output"))
    if data_analyst:
        section_num += 1
        doc.add_heading(f"{section_num}. Data Analysis", level=1)
        da_narrative = data_analyst.get("narrative", "")
        if da_narrative:
            _add_narrative_paragraphs(doc, da_narrative)

        if isinstance(data_analyst, dict):
            analysis = data_analyst.get("data_analysis", {})
            if isinstance(analysis, dict):
                for k, v in analysis.items():
                    if isinstance(v, str) and len(v) > 20:
                        doc.add_heading(k.replace("_", " ").title(), level=2)
                        _add_narrative_paragraphs(doc, v)
                    elif isinstance(v, list) and v:
                        doc.add_heading(k.replace("_", " ").title(), level=2)
                        for item in v[:15]:
                            doc.add_paragraph(str(item), style="List Bullet")
        doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # FEEDBACK & LEARNING OUTPUT (if available)
    # ══════════════════════════════════════════════════════
    feedback = _unwrap_agent_output(bid_data.get("feedback_output"))
    if feedback:
        section_num += 1
        doc.add_heading(f"{section_num}. Feedback & Quality Insights", level=1)
        fb_narrative = feedback.get("narrative", "")
        if fb_narrative:
            _add_narrative_paragraphs(doc, fb_narrative)

        pipeline_health = feedback.get("pipeline_health_score")
        if pipeline_health:
            doc.add_paragraph(f"Pipeline Health Score: {pipeline_health}/100")

        contradictions = feedback.get("contradictions", [])
        if contradictions:
            doc.add_heading("Cross-Agent Contradictions", level=2)
            for c in contradictions:
                if isinstance(c, dict):
                    doc.add_paragraph(
                        f"{c.get('agents', '')}: {c.get('description', c.get('issue', ''))}",
                        style="List Bullet",
                    )
        doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # DYNAMIC CATCH-ALL: Any agent outputs not explicitly handled above
    # ══════════════════════════════════════════════════════
    _HANDLED_KEYS = {
        "intake_output",
        "scope_output",
        "solution_output",
        "competitive_output",
        "commercial_output",
        "compliance_output",
        "automation_ai_output",
        "transition_change_output",
        "client_intel_output",
        "proposal_output",
        "discovery_output",
        "qa_output",
        "output_generator_output",
        "data_analyst_output",
        "feedback_output",
    }
    _AGENT_LABELS = {
        "intake_output": "Intake",
        "data_analyst_output": "Data Analysis",
        "client_intel_output": "Client Intelligence",
        "bid_no_bid_output": "Bid/No-Bid",
        "scope_output": "Scope",
        "solution_output": "Solution Architecture",
        "automation_ai_output": "Automation & AI",
        "competitive_output": "Competitive Intel",
        "commercial_output": "Commercial Model",
        "compliance_output": "Compliance & Risk",
        "proposal_output": "Proposal",
        "output_generator_output": "Output Generator",
        "discovery_output": "Discovery",
        "qa_output": "QA",
        "feedback_output": "Feedback",
        "transition_change_output": "Transition & Change",
    }

    for key in list(bid_data.keys()):
        if not key.endswith("_output") or key in _HANDLED_KEYS:
            continue
        agent_data = _unwrap_agent_output(bid_data.get(key))
        if not agent_data or not isinstance(agent_data, dict):
            continue

        label = _AGENT_LABELS.get(
            key, key.replace("_output", "").replace("_", " ").title()
        )
        section_num += 1
        doc.add_heading(f"{section_num}. {label}", level=1)

        # Try narrative first
        narrative = agent_data.get("narrative", "")
        if narrative:
            _add_narrative_paragraphs(doc, narrative)

        # Render remaining dict content
        for sub_key, sub_val in agent_data.items():
            if sub_key == "narrative":
                continue
            if isinstance(sub_val, str) and len(sub_val) > 30:
                doc.add_heading(sub_key.replace("_", " ").title(), level=2)
                _add_narrative_paragraphs(doc, sub_val)
            elif isinstance(sub_val, list) and sub_val:
                doc.add_heading(sub_key.replace("_", " ").title(), level=2)
                if isinstance(sub_val[0], dict):
                    headers = list(sub_val[0].keys())[:5]
                    rows = [
                        [str(item.get(h, ""))[:60] for h in headers]
                        for item in sub_val[:15]
                    ]
                    _add_table(
                        doc, [h.replace("_", " ").title() for h in headers], rows
                    )
                else:
                    for item in sub_val[:15]:
                        doc.add_paragraph(str(item), style="List Bullet")
            elif isinstance(sub_val, dict):
                doc.add_heading(sub_key.replace("_", " ").title(), level=2)
                for k, v in sub_val.items():
                    doc.add_paragraph(
                        f"{k.replace('_', ' ').title()}: {str(v)[:200]}",
                        style="List Bullet",
                    )

        doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # CLOSING
    # ══════════════════════════════════════════════════════
    section_num += 1
    doc.add_heading(f"{section_num}. Next Steps", level=1)
    bidder_name = (
        bid_data.get("manifest", {}).get("bidder_profile", {}).get("name", "[Bidder]")
    )
    doc.add_paragraph(
        f"{bidder_name} looks forward to partnering with {client_name} on this engagement. "
        f"We are confident in our ability to deliver exceptional value through our proven "
        f"delivery methodology and deep domain expertise."
    )
    doc.add_paragraph(
        f"For questions or clarifications, please contact your designated {bidder_name} engagement manager."
    )
    doc.add_paragraph("")
    doc.add_paragraph(f"Document Reference: {bid_ref}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}")

    # Save
    os.makedirs(output_dir, exist_ok=True)
    fname = (
        filename
        or f"BidResponse_{client_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
    )
    path = os.path.join(output_dir, fname)
    doc.save(path)

    return {
        "status": "success",
        "path": path,
        "filename": fname,
        "pages_estimate": section_num * 3 + 5,
    }
